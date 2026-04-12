"""Export services for manuscripts."""
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
import subprocess
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from utils.config import VAULT_ROOT

# Lazy imports for optional dependencies
try:
    import ebooklib
    from ebooklib import epub
    HAS_EBOOKLIB = True
except ImportError:
    HAS_EBOOKLIB = False

# Check if pandoc is available
def _check_pandoc() -> bool:
    """Check if pandoc is installed and available."""
    return shutil.which("pandoc") is not None


def export_docx_manuscript(
    book_id: str,
    manuscript_text: str,
    working_title: str,
    author: Optional[str] = None,
    date: Optional[str] = None
) -> Path:
    """
    Export manuscript as DOCX in Publisher Standard Manuscript format.
    
    Format rules:
    - 12pt Times New Roman (fallback serif)
    - Double spaced
    - First-line indent 0.5"
    - Page numbers
    - Title page (working_title, author, date)
    
    Returns path to exported file.
    """
    # Create exports directory
    exports_dir = VAULT_ROOT / "projects" / book_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    
    # Create document
    doc = Document()
    
    # Set default font to Times New Roman, 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph format: double spaced, first-line indent 0.5"
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0  # Double spaced
    paragraph_format.first_line_indent = Inches(0.5)
    paragraph_format.space_after = Pt(0)
    
    # Add title page
    title_page = doc.add_paragraph()
    title_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title_run = title_page.add_run(working_title)
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    # Spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Author (if available)
    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author)
        author_run.font.size = Pt(14)
    
    # Date (if available, otherwise use current date)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = date or datetime.now().strftime("%B %d, %Y")
    date_run = date_para.add_run(date_str)
    date_run.font.size = Pt(12)
    
    # Page break after title page
    doc.add_page_break()
    
    # Process manuscript text
    # Split by sections (assuming markdown-style headers)
    lines = manuscript_text.split('\n')
    current_paragraph = None
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines (but add spacing between sections)
        if not line:
            if current_paragraph:
                doc.add_paragraph()  # Add spacing
            current_paragraph = None
            continue
        
        # Check if it's a header (starts with #)
        if line.startswith('#'):
            # Remove # and create heading
            heading_text = line.lstrip('#').strip()
            if heading_text:
                heading = doc.add_heading(heading_text, level=1)
                heading_format = heading.paragraph_format
                heading_format.line_spacing = 2.0
                heading_format.space_after = Pt(12)
                current_paragraph = None
        else:
            # Regular paragraph
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.line_spacing = 2.0
            para_format.first_line_indent = Inches(0.5)
            para_format.space_after = Pt(0)
            
            # Add text
            para.add_run(line)
            current_paragraph = para
    
    # Add page numbers to footer
    # Note: python-docx doesn't have a high-level API for page number fields
    # We'll add a simple footer that Word will recognize and update
    section = doc.sections[0]
    footer = section.footer
    # Clear existing paragraphs
    for para in footer.paragraphs:
        para.clear()
    footer_para = footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run()
    footer_run.font.size = Pt(10)
    footer_run.font.name = 'Times New Roman'
    
    # Add page number using XML (proper way)
    from docx.oxml import OxmlElement
    
    # Add "Page " text
    footer_run.add_text("Page ")
    
    # Create page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_run._element.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    footer_run._element.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run._element.append(fldChar2)
    
    # Save file
    filename = f"{working_title.replace(' ', '_')}_manuscript.docx"
    filepath = exports_dir / filename
    doc.save(str(filepath))
    
    return filepath


def export_epub_manuscript(
    book_id: str,
    manuscript_text: str,
    working_title: str,
    author: Optional[str] = None,
    date: Optional[str] = None,
    language: str = "en"
) -> Path:
    """
    Export manuscript as EPUB.
    
    Prefers pandoc if installed, falls back to ebooklib.
    Includes metadata: title, author, language, date.
    Cover placeholder is optional.
    
    Returns path to exported file.
    """
    # Create exports directory
    exports_dir = VAULT_ROOT / "projects" / book_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    
    filename = f"{working_title.replace(' ', '_')}_manuscript.epub"
    filepath = exports_dir / filename
    
    # Try pandoc first (preferred method)
    if _check_pandoc():
        return _export_epub_with_pandoc(
            manuscript_text=manuscript_text,
            working_title=working_title,
            author=author,
            date=date,
            language=language,
            filepath=filepath
        )
    elif HAS_EBOOKLIB:
        # Fallback to ebooklib
        return _export_epub_with_ebooklib(
            manuscript_text=manuscript_text,
            working_title=working_title,
            author=author,
            date=date,
            language=language,
            filepath=filepath
        )
    else:
        raise ImportError(
            "Neither pandoc nor ebooklib is available. "
            "Install pandoc (https://pandoc.org/installing.html) or "
            "ebooklib (pip install ebooklib)"
        )


def _export_epub_with_pandoc(
    manuscript_text: str,
    working_title: str,
    author: Optional[str],
    date: Optional[str],
    language: str,
    filepath: Path
) -> Path:
    """Export EPUB using pandoc (preferred method)."""
    # Create temporary markdown file
    import tempfile
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
        tmp.write(f"# {working_title}\n\n")
        if author:
            tmp.write(f"**Author:** {author}\n\n")
        if date:
            tmp.write(f"**Date:** {date}\n\n")
        tmp.write("---\n\n")
        tmp.write(manuscript_text)
        tmp_path = tmp.name
    
    try:
        # Build pandoc command
        cmd = [
            "pandoc",
            tmp_path,
            "-o", str(filepath),
            "--epub-metadata", f"title={working_title}",
            "--epub-cover-image", "none",  # Placeholder, can be enhanced later
        ]
        
        if author:
            cmd.extend(["--epub-metadata", f"creator={author}"])
        
        if language:
            cmd.extend(["--epub-metadata", f"language={language}"])
        
        # Run pandoc
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            check=True
        )
        
        return filepath
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Pandoc failed: {e.stderr}")
    finally:
        # Clean up temp file
        Path(tmp_path).unlink(missing_ok=True)


def _export_epub_with_ebooklib(
    manuscript_text: str,
    working_title: str,
    author: Optional[str],
    date: Optional[str],
    language: str,
    filepath: Path
) -> Path:
    """Export EPUB using ebooklib (fallback method)."""
    book = epub.EpubBook()
    
    # Set metadata
    book.set_identifier(f"book_{datetime.now().timestamp()}")
    book.set_title(working_title)
    book.set_language(language)
    
    if author:
        book.add_author(author)
    
    if date:
        book.add_metadata("DC", "date", date)
    
    # Process manuscript text into chapters
    # Split by markdown headers (#) to create chapters
    lines = manuscript_text.split('\n')
    chapters = []
    current_chapter = []
    current_title = None
    
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            # Save previous chapter if exists
            if current_chapter and current_title:
                chapters.append((current_title, '\n'.join(current_chapter)))
            # Start new chapter
            current_title = line.lstrip('#').strip()
            current_chapter = []
        else:
            if line or current_chapter:  # Keep empty lines if we have content
                current_chapter.append(line)
    
    # Add final chapter
    if current_chapter and current_title:
        chapters.append((current_title, '\n'.join(current_chapter)))
    
    # If no chapters found, create one with all content
    if not chapters:
        chapters = [("Chapter 1", manuscript_text)]
    
    # Create chapters
    spine = []
    toc = []
    
    for i, (title, content) in enumerate(chapters, 1):
        chapter_id = f"chapter_{i}"
        chapter = epub.EpubHtml(
            title=title,
            file_name=f"{chapter_id}.xhtml",
            lang=language
        )
        # Convert markdown-style headers to HTML
        html_content = _markdown_to_html(content)
        chapter.content = f'<h1>{title}</h1>{html_content}'
        
        book.add_item(chapter)
        spine.append(chapter)
        toc.append(chapter)
    
    # Add default NCX and Nav file
    book.toc = toc
    book.spine = ['nav'] + spine
    
    # Add navigation files
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Write the EPUB file
    epub.write_epub(str(filepath), book)
    
    return filepath


def _markdown_to_html(text: str) -> str:
    """Simple markdown to HTML converter for basic formatting."""
    lines = text.split('\n')
    html_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            html_lines.append('<p></p>')
            continue
        
        # Headers
        if line.startswith('##'):
            content = line.lstrip('#').strip()
            html_lines.append(f'<h2>{content}</h2>')
        elif line.startswith('#'):
            content = line.lstrip('#').strip()
            html_lines.append(f'<h1>{content}</h1>')
        else:
            # Basic paragraph (escape HTML)
            escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            html_lines.append(f'<p>{escaped}</p>')
    
    return '\n'.join(html_lines)

