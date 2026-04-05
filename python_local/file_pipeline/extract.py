"""Text extraction — pulls text from PDF, DOCX, plain text, and images."""

from __future__ import annotations

import logging
from pathlib import Path

from .config import OCR_EXTENSIONS

log = logging.getLogger(__name__)


def extract_text(path: Path) -> dict:
    """Extract text from a file. Returns dict with 'text', 'method', 'pages'."""
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(path)
    if ext in (".txt", ".md", ".csv"):
        return _extract_plain(path)
    if ext in (".docx",):
        return _extract_docx(path)
    if ext in OCR_EXTENSIONS:
        return _extract_image_ocr(path)

    log.warning("Unsupported extraction format: %s", ext)
    return {"text": "", "method": "unsupported", "pages": 0}


def _extract_pdf(path: Path) -> dict:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfStreamError, PdfReadError
    except ImportError:
        log.error("pypdf not installed — cannot extract PDF")
        return {"text": "", "method": "missing_dependency", "pages": 0}

    try:
        reader = PdfReader(str(path))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages_text.append(text)

        full_text = "\n\n".join(pages_text)
        return {"text": full_text, "method": "pypdf", "pages": len(reader.pages)}
    except (PdfStreamError, PdfReadError) as exc:
        log.warning("Corrupt PDF %s: %s", path.name, exc)
        return {"text": "", "method": "pypdf_error", "pages": 0}


def _extract_plain(path: Path) -> dict:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        log.error("Failed to read %s: %s", path.name, exc)
        return {"text": "", "method": "read_error", "pages": 0}
    return {"text": text, "method": "plain_text", "pages": 1}


def _extract_docx(path: Path) -> dict:
    try:
        import docx
    except ImportError:
        log.error("python-docx not installed — cannot extract DOCX")
        return {"text": "", "method": "missing_dependency", "pages": 0}

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return {"text": text, "method": "python-docx", "pages": len(doc.paragraphs)}


def _extract_image_ocr(path: Path) -> dict:
    """Extract text from images using Tesseract OCR (if available)."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        log.warning("pytesseract/Pillow not installed — skipping OCR for %s", path.name)
        return {"text": "", "method": "ocr_unavailable", "pages": 1}

    try:
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return {"text": text, "method": "tesseract_ocr", "pages": 1}
    except Exception as exc:
        log.error("OCR failed for %s: %s", path.name, exc)
        return {"text": "", "method": "ocr_error", "pages": 1}
